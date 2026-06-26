from __future__ import annotations

import hashlib
import html
import io
import json
import os
import re
import tempfile
import uuid
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Mapping

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from app.utils.offline_engines import extract_sections, normalize_text, skill_extraction_engine


PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_PDF_PAGES = 100
SAFE_SECTION_NAME_RE = re.compile(r"[^a-z0-9_]+")


class ResumeFileValidationError(ValueError):
    def __init__(self, message: str, field: str = "file") -> None:
        super().__init__(message)
        self.field = field


ResumeValidationError = ResumeFileValidationError


class ResumeParsingError(RuntimeError):
    """Raised when an uploaded resume cannot be parsed safely."""


class ResumeExportError(RuntimeError):
    """Raised when a resume export cannot be generated."""


@dataclass(frozen=True)
class ValidatedResumeUpload:
    original_filename: str
    secure_filename: str
    extension: str
    mime_type: str
    size_bytes: int
    checksum_sha256: str
    content: bytes


@dataclass(frozen=True)
class ParsedResume:
    text: str
    sections: dict[str, str]
    skills: tuple[str, ...]
    skills_by_category: dict[str, tuple[str, ...]]
    parser: str
    page_count: int | None
    confidence_score: float

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @property
    def confidence(self) -> float:
        return self.confidence_score


@dataclass(frozen=True)
class ExportPayload:
    content: bytes
    mime_type: str
    extension: str


def validate_resume_upload(
    file_storage: FileStorage | None,
    *,
    max_bytes: int,
    allowed_extensions: set[str] | frozenset[str],
) -> ValidatedResumeUpload:
    if file_storage is None or not file_storage.filename:
        raise ResumeFileValidationError("Select a PDF or DOCX resume to upload.")

    original_filename = file_storage.filename.strip()
    safe_name = secure_filename(original_filename)
    if not safe_name or "." not in safe_name:
        raise ResumeFileValidationError("The uploaded filename is invalid.")

    extension = safe_name.rsplit(".", 1)[1].lower()
    if extension not in allowed_extensions:
        raise ResumeFileValidationError("Only PDF and DOCX files are allowed.")

    content = file_storage.stream.read(max_bytes + 1)
    file_storage.stream.seek(0)
    if not content:
        raise ResumeFileValidationError("The uploaded file is empty.")
    if len(content) > max_bytes:
        raise ResumeFileValidationError("The uploaded file exceeds the 5MB limit.")

    detected_extension, mime_type = detect_resume_file_type(content)
    if detected_extension != extension:
        raise ResumeFileValidationError("The file content does not match its extension.")

    return ValidatedResumeUpload(
        original_filename=original_filename,
        secure_filename=safe_name,
        extension=extension,
        mime_type=mime_type,
        size_bytes=len(content),
        checksum_sha256=hashlib.sha256(content).hexdigest(),
        content=content,
    )


def detect_resume_file_type(content: bytes) -> tuple[str, str]:
    if content.startswith(b"%PDF-"):
        if b"%%EOF" not in content[-4096:]:
            raise ResumeFileValidationError("The PDF file appears incomplete or corrupted.")
        return "pdf", PDF_MIME_TYPE

    if content.startswith(b"PK\x03\x04"):
        _validate_docx_container(content)
        return "docx", DOCX_MIME_TYPE

    raise ResumeFileValidationError("The file signature is not a valid PDF or DOCX document.")


def store_validated_upload(upload: ValidatedResumeUpload, storage_root: str, user_id: int) -> Path:
    user_root = _safe_user_directory(storage_root, user_id)
    stored_name = f"{uuid.uuid4().hex}_{upload.secure_filename}"
    destination = _safe_child_path(user_root, stored_name)
    _atomic_write(destination, upload.content)
    return destination


def save_validated_upload(upload: ValidatedResumeUpload, storage_root: str, user_id: int) -> str:
    return str(store_validated_upload(upload, storage_root, user_id))


def store_export(content: bytes, storage_root: str, user_id: int, basename: str, extension: str) -> Path:
    user_root = _safe_user_directory(storage_root, user_id)
    safe_base = secure_filename(basename) or "resume"
    destination = _safe_child_path(user_root, f"{safe_base}_{uuid.uuid4().hex[:12]}.{extension}")
    _atomic_write(destination, content)
    return destination


def parse_resume_file(path: str | Path, mime_type: str) -> ParsedResume:
    file_path = Path(path).resolve()
    if not file_path.is_file():
        raise ResumeParsingError("The stored resume file could not be found.")

    if mime_type == PDF_MIME_TYPE:
        text, parser_name, page_count = _parse_pdf(file_path)
    elif mime_type == DOCX_MIME_TYPE:
        text, parser_name, page_count = _parse_docx(file_path)
    else:
        raise ResumeParsingError("Unsupported resume file type.")

    normalized = normalize_text(text)
    if len(normalized) < 20:
        raise ResumeParsingError("No readable resume text could be extracted from this file.")

    sections = extract_sections(normalized)
    skill_result = skill_extraction_engine.extract(normalized)
    confidence = _parsing_confidence(normalized, sections, skill_result.skills, page_count)
    return ParsedResume(
        text=normalized,
        sections=sections,
        skills=skill_result.skills,
        skills_by_category=skill_result.skills_by_category,
        parser=parser_name,
        page_count=page_count,
        confidence_score=confidence,
    )


def content_to_plain_text(content: Mapping[str, Any] | None) -> str:
    if not content:
        return ""
    lines: list[str] = []
    for section_name, section_value in content.items():
        title = str(section_name).replace("_", " ").strip().title()
        section_text = _value_to_text(section_value)
        if section_text:
            lines.extend((title, section_text, ""))
    return normalize_text("\n".join(lines))


def plain_text_from_content(content: Mapping[str, Any] | None) -> str:
    return content_to_plain_text(content)


def normalize_builder_content(content: Mapping[str, Any]) -> dict[str, Any]:
    normalized: dict[str, Any] = {}
    for raw_key, value in content.items():
        key = SAFE_SECTION_NAME_RE.sub("_", str(raw_key).strip().lower()).strip("_")
        if not key or len(key) > 64:
            continue
        normalized[key] = _normalize_json_value(value)
    if not normalized:
        raise ValueError("Resume content must include at least one section.")
    return normalized


def content_from_sections(sections: Mapping[str, Any]) -> dict[str, Any]:
    return normalize_builder_content(sections)


def export_resume_payload(
    export_format: str,
    *,
    title: str,
    content: Mapping[str, Any],
    metadata: Mapping[str, Any] | None = None,
) -> ExportPayload:
    normalized_format = export_format.strip().lower()
    if normalized_format == "json":
        return ExportPayload(_export_json(title, content, metadata), "application/json", "json")
    if normalized_format == "html":
        return ExportPayload(_export_html(title, content), "text/html; charset=utf-8", "html")
    if normalized_format == "docx":
        return ExportPayload(
            _export_docx(title, content),
            DOCX_MIME_TYPE,
            "docx",
        )
    if normalized_format == "pdf":
        return ExportPayload(_export_pdf(title, content), PDF_MIME_TYPE, "pdf")
    raise ResumeExportError("Supported export formats are PDF, JSON, HTML, and DOCX.")


def build_json_export(payload: Mapping[str, Any]) -> bytes:
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def build_printable_html(title: str, content: Mapping[str, Any]) -> str:
    return _export_html(title, content).decode("utf-8")


def build_docx_export(title: str, content: Mapping[str, Any]) -> bytes:
    return _export_docx(title, content)


def build_pdf_export(title: str, content: Mapping[str, Any]) -> bytes:
    return _export_pdf(title, content)


def _validate_docx_container(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            names = {member.filename for member in members}
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise ResumeFileValidationError("The uploaded ZIP is not a valid DOCX document.")

            total_size = 0
            for member in members:
                path = Path(member.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise ResumeFileValidationError("The DOCX archive contains unsafe paths.")
                total_size += member.file_size
                if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ResumeFileValidationError("The DOCX archive expands beyond the safe processing limit.")
                if member.compress_size and member.file_size / member.compress_size > 200:
                    raise ResumeFileValidationError("The DOCX archive has an unsafe compression ratio.")
    except zipfile.BadZipFile as exc:
        raise ResumeFileValidationError("The DOCX file is corrupted.") from exc


def _parse_pdf(path: Path) -> tuple[str, str, int]:
    errors: list[str] = []
    try:
        import pdfplumber

        with pdfplumber.open(path) as document:
            if len(document.pages) > MAX_PDF_PAGES:
                raise ResumeParsingError(f"PDF exceeds the {MAX_PDF_PAGES}-page processing limit.")
            pages = [(page.extract_text(x_tolerance=2, y_tolerance=3) or "") for page in document.pages]
            text = "\n\n".join(pages)
            if len(normalize_text(text)) >= 20:
                return text, "pdfplumber", len(document.pages)
    except ResumeParsingError:
        raise
    except Exception as exc:
        errors.append(f"pdfplumber:{exc.__class__.__name__}")

    try:
        import fitz

        with fitz.open(path) as document:
            if document.page_count > MAX_PDF_PAGES:
                raise ResumeParsingError(f"PDF exceeds the {MAX_PDF_PAGES}-page processing limit.")
            pages = [page.get_text("text", sort=True) or "" for page in document]
            text = "\n\n".join(pages)
            if len(normalize_text(text)) >= 20:
                return text, "pymupdf", document.page_count
    except ResumeParsingError:
        raise
    except Exception as exc:
        errors.append(f"pymupdf:{exc.__class__.__name__}")

    detail = ", ".join(errors) if errors else "no readable text"
    raise ResumeParsingError(f"PDF text extraction failed ({detail}).")


def _parse_docx(path: Path) -> tuple[str, str, None]:
    try:
        from docx import Document

        document = Document(path)
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts), "python-docx", None
    except Exception as exc:
        raise ResumeParsingError("DOCX text extraction failed.") from exc


def _parsing_confidence(text: str, sections: Mapping[str, str], skills: tuple[str, ...], page_count: int | None) -> float:
    word_count = len(text.split())
    section_score = min(len(sections), 7) * 8
    skill_score = min(len(skills), 12) * 2.5
    length_score = 30 if 250 <= word_count <= 1200 else 20 if word_count >= 100 else 8
    page_score = 10 if page_count is None or 1 <= page_count <= 4 else 5
    return round(min(section_score + skill_score + length_score + page_score, 100), 2)


def _safe_user_directory(root: str, user_id: int) -> Path:
    root_path = Path(root).resolve()
    root_path.mkdir(parents=True, exist_ok=True)
    user_root = (root_path / str(int(user_id))).resolve()
    if root_path not in user_root.parents:
        raise ResumeFileValidationError("Unsafe storage path.")
    user_root.mkdir(parents=True, exist_ok=True)
    return user_root


def _safe_child_path(parent: Path, filename: str) -> Path:
    destination = (parent / filename).resolve()
    if parent != destination.parent:
        raise ResumeFileValidationError("Unsafe storage filename.")
    return destination


def _atomic_write(destination: Path, content: bytes) -> None:
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(dir=destination.parent, delete=False) as temporary:
            temporary.write(content)
            temporary.flush()
            os.fsync(temporary.fileno())
            temp_path = Path(temporary.name)
        os.replace(temp_path, destination)
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)


def _normalize_json_value(value: Any) -> Any:
    if value is None or isinstance(value, (bool, int, float)):
        return value
    if isinstance(value, str):
        return normalize_text(value)
    if isinstance(value, Mapping):
        return {str(key)[:80]: _normalize_json_value(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_normalize_json_value(item) for item in value]
    return normalize_text(str(value))


def _value_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return normalize_text(value)
    if isinstance(value, Mapping):
        return normalize_text("\n".join(_value_to_text(item) for item in value.values()))
    if isinstance(value, (list, tuple)):
        return normalize_text("\n".join(_value_to_text(item) for item in value))
    return normalize_text(str(value))


def _export_json(title: str, content: Mapping[str, Any], metadata: Mapping[str, Any] | None) -> bytes:
    payload = {"title": title, "content": content, "metadata": dict(metadata or {})}
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def _export_html(title: str, content: Mapping[str, Any]) -> bytes:
    sections = []
    for name, value in content.items():
        section_title = html.escape(str(name).replace("_", " ").title())
        body = html.escape(_value_to_text(value)).replace("\n", "<br>")
        sections.append(f"<section><h2>{section_title}</h2><p>{body}</p></section>")
    document = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{ color: #17202a; font: 11pt/1.5 Arial, sans-serif; margin: 0 auto; max-width: 820px; padding: 36px; }}
    h1 {{ border-bottom: 2px solid #1f6f78; font-size: 24pt; margin: 0 0 24px; padding-bottom: 10px; }}
    h2 {{ color: #1f6f78; font-size: 13pt; margin: 22px 0 8px; text-transform: uppercase; }}
    p {{ margin: 0; white-space: normal; }}
    @media print {{ body {{ max-width: none; padding: 0; }} }}
  </style>
</head>
<body><h1>{html.escape(title)}</h1>{''.join(sections)}</body>
</html>"""
    return document.encode("utf-8")


def _export_docx(title: str, content: Mapping[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt

        document = Document()
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = Pt(10.5)
        document.add_heading(title, level=0)
        for name, value in content.items():
            document.add_heading(str(name).replace("_", " ").title(), level=1)
            for line in _value_to_text(value).splitlines():
                clean = line.strip()
                if clean:
                    document.add_paragraph(clean)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:
        raise ResumeExportError("DOCX export failed.") from exc


def _export_pdf(title: str, content: Mapping[str, Any]) -> bytes:
    try:
        from fpdf import FPDF

        pdf = FPDF(format="A4")
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_title(_pdf_safe(title))
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 9, _pdf_safe(title))
        pdf.ln(2)
        for name, value in content.items():
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(31, 111, 120)
            pdf.multi_cell(0, 7, _pdf_safe(str(name).replace("_", " ").upper()))
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(23, 32, 42)
            pdf.multi_cell(0, 5.5, _pdf_safe(_value_to_text(value)))
            pdf.ln(2)
        return bytes(pdf.output())
    except Exception as exc:
        raise ResumeExportError("PDF export failed.") from exc


def _pdf_safe(value: str) -> str:
    return value.encode("latin-1", errors="replace").decode("latin-1")
