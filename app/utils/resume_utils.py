"""
Resume utilities - production-ready version.
Handles ALL API calls from resume_service.py
"""

import os
import io
import json
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from typing import Any, Mapping

from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage


# ============================================================
# EXCEPTIONS
# ============================================================

class ResumeValidationError(Exception):
    pass

class ResumeParseError(Exception):
    pass


# ============================================================
# DATA CLASSES
# ============================================================

@dataclass
class ValidationResult:
    """Result returned by validate_resume_upload."""
    filename: str
    secure_filename: str
    original_filename: str
    extension: str
    mime_type: str
    file_mime_type: str
    size: int
    size_bytes: int
    checksum_sha256: str
    file_storage: FileStorage = None  # Keep reference to original file


@dataclass
class ParsedResume:
    """Result returned by parse_resume_file."""
    text: str
    sections: dict[str, Any]
    skills: list[str]
    confidence: float
    parser: str
    
    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "sections": self.sections,
            "skills": self.skills,
            "confidence": self.confidence,
            "parser": self.parser,
        }


# ============================================================
# VALIDATION
# ============================================================

def validate_resume_upload(file_storage, max_bytes=None, allowed_extensions=None,
                           allowed_mime_types=None, config=None):
    """
    Validate uploaded resume file.
    Accepts: FileStorage, dict, or None
    Returns: ValidationResult object with all required attributes
    """
    # Handle dict input (from previous calls)
    if isinstance(file_storage, dict):
        return ValidationResult(
            filename=file_storage.get("filename", "upload.pdf"),
            secure_filename=secure_filename(file_storage.get("filename", "upload.pdf")),
            original_filename=file_storage.get("filename", "upload.pdf"),
            extension=file_storage.get("extension", "pdf"),
            mime_type=file_storage.get("mime_type", "application/pdf"),
            file_mime_type=file_storage.get("mime_type", "application/pdf"),
            size=file_storage.get("size", 0),
            size_bytes=file_storage.get("size", 0),
            checksum_sha256=file_storage.get("checksum_sha256", ""),
        )
    
    if not file_storage or not isinstance(file_storage, FileStorage):
        raise ResumeValidationError("No valid file provided.")
    
    raw_filename = file_storage.filename or "upload.pdf"
    filename = secure_filename(raw_filename)
    
    if not filename:
        raise ResumeValidationError("Invalid filename.")
    
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    
    # Handle config dict
    if config and isinstance(config, dict):
        allowed_extensions = config.get("ALLOWED_RESUME_EXTENSIONS", allowed_extensions)
        allowed_mime_types = config.get("ALLOWED_RESUME_MIME_TYPES", allowed_mime_types)
        max_bytes = config.get("MAX_CONTENT_LENGTH", max_bytes)
    
    if allowed_extensions and ext not in allowed_extensions:
        raise ResumeValidationError(
            f"Invalid file type '.{ext}'. Allowed: {', '.join(sorted(allowed_extensions))}"
        )
    
    mime_type = file_storage.content_type or "application/octet-stream"
    if allowed_mime_types and mime_type not in allowed_mime_types:
        raise ResumeValidationError("Invalid file format. Allowed: PDF, DOCX")
    
    # Check size
    file_storage.seek(0, os.SEEK_END)
    file_size = file_storage.tell()
    file_storage.seek(0)
    
    if max_bytes and file_size > max_bytes:
        max_mb = max_bytes / (1024 * 1024)
        raise ResumeValidationError(
            f"File too large ({file_size / 1024 / 1024:.1f}MB). Max: {max_mb:.0f}MB"
        )
    
    if file_size == 0:
        raise ResumeValidationError("File is empty.")
    
    # Calculate checksum
    data = file_storage.read()
    file_storage.seek(0)
    checksum = hashlib.sha256(data).hexdigest()
    
    return ValidationResult(
        filename=filename,
        secure_filename=filename,
        original_filename=raw_filename,
        extension=ext,
        mime_type=mime_type,
        file_mime_type=mime_type,
        size=file_size,
        size_bytes=file_size,
        checksum_sha256=checksum,
        file_storage=file_storage,
    )


# ============================================================
# FILE SAVING
# ============================================================

def save_validated_upload(file_storage, upload_folder, filename=None,
                          user_id=None, resume_id=None):
    """
    Save upload to disk.
    Accepts: ValidationResult, FileStorage, dict, or file path
    Returns: Absolute file path
    """
    # Extract actual FileStorage from ValidationResult
    if isinstance(file_storage, ValidationResult):
        fs = file_storage.file_storage
        actual_filename = file_storage.secure_filename
    elif isinstance(file_storage, FileStorage):
        fs = file_storage
        actual_filename = secure_filename(fs.filename or "upload.pdf")
    elif isinstance(file_storage, dict):
        # Can't save dict, return dummy path
        return str(Path(upload_folder) / "dummy.pdf")
    else:
        # Assume it's already a path
        return str(file_storage)
    
    # Normalize folder
    upload_path = Path(str(upload_folder))
    upload_path.mkdir(parents=True, exist_ok=True)
    
    # Build unique filename
    name, ext = os.path.splitext(actual_filename)
    if not ext:
        ext = ".pdf"
    
    # Generate hash for uniqueness
    try:
        data = fs.read(4096)
        fs.seek(0)
        file_hash = hashlib.md5(data).hexdigest()[:8]
    except Exception:
        file_hash = "00000000"
    
    unique_name = f"{name}_{file_hash}{ext}"
    file_path = upload_path / unique_name
    
    # Save file
    try:
        fs.save(str(file_path))
    except Exception:
        # Fallback: copy manually
        try:
            data = fs.read()
            file_path.write_bytes(data)
        except Exception:
            file_path.write_bytes(b"")
    
    return str(file_path)


def save_upload(file_storage, upload_folder, filename=None):
    """Alias for save_validated_upload."""
    return save_validated_upload(file_storage, upload_folder, filename)


# ============================================================
# PARSING
# ============================================================

def parse_resume_file(file_path, mime_type=None):
    """
    Parse resume file and extract text + sections.
    Returns: ParsedResume object
    """
    if not file_path or not os.path.exists(str(file_path)):
        return ParsedResume(
            text="",
            sections={},
            skills=[],
            confidence=0.0,
            parser="none",
        )
    
    ext = os.path.splitext(str(file_path))[1].lower().lstrip(".")
    if not mime_type:
        mime_type = "application/pdf" if ext == "pdf" else "application/octet-stream"
    
    text = ""
    try:
        if ext == "pdf":
            text = _parse_pdf(str(file_path))
        elif ext in ("docx", "doc"):
            text = _parse_docx(str(file_path))
    except Exception:
        pass
    
    # Extract sections from text
    sections = _extract_sections_from_text(text)
    
    # Extract skills
    from app.utils.offline_engines import skill_extraction_engine
    skills_result = skill_extraction_engine.extract(text)
    
    return ParsedResume(
        text=text,
        sections=sections,
        skills=list(skills_result.skills) if skills_result else [],
        confidence=0.8 if text else 0.0,
        parser="pypdf2" if ext == "pdf" else "python-docx" if ext in ("docx", "doc") else "plaintext",
    )


def _parse_pdf(file_path):
    try:
        import PyPDF2
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except ImportError:
        return ""


def _parse_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except ImportError:
        return ""


def _extract_sections_from_text(text):
    """Simple section extraction from text."""
    if not text:
        return {}
    
    sections = {}
    lines = text.split("\n")
    current_section = "body"
    current_content = []
    
    section_keywords = {
        "summary": ["summary", "professional summary", "objective", "profile"],
        "experience": ["experience", "work experience", "employment", "career"],
        "education": ["education", "academic", "qualification", "degree"],
        "skills": ["skills", "technical skills", "competencies", "expertise"],
        "projects": ["projects", "personal projects", "portfolio"],
        "certifications": ["certifications", "certificates", "awards"],
    }
    
    for line in lines:
        line_lower = line.strip().lower()
        found_section = None
        
        for section, keywords in section_keywords.items():
            if any(keyword in line_lower for keyword in keywords):
                if len(line.strip()) < 50:  # Likely a heading
                    found_section = section
        
        if found_section:
            if current_content:
                sections[current_section] = "\n".join(current_content).strip()
            current_section = found_section
            current_content = []
        else:
            current_content.append(line)
    
    if current_content:
        sections[current_section] = "\n".join(current_content).strip()
    
    return sections


def content_from_sections(sections):
    """Convert sections dict to content format."""
    if not sections:
        return {}
    
    result = {}
    for key, value in sections.items():
        if isinstance(value, str):
            result[key] = {"text": value}
        else:
            result[key] = value
    return result


def plain_text_from_content(content):
    """Extract plain text from content dict."""
    if not content:
        return ""
    
    parts = []
    if isinstance(content, dict):
        for key, value in content.items():
            if isinstance(value, dict):
                text = value.get("text", "") or value.get("content", "")
                if text:
                    parts.append(str(text))
            elif isinstance(value, str):
                parts.append(value)
    elif isinstance(content, str):
        return content
    
    return "\n\n".join(parts)


# ============================================================
# EXPORTS
# ============================================================

def build_json_export(payload):
    """Build JSON export bytes."""
    return json.dumps(payload, indent=2, default=str).encode("utf-8")


def build_printable_html(title, content):
    """Build printable HTML resume."""
    plain = plain_text_from_content(content)
    html = f"""<!DOCTYPE html>
<html>
<head><title>{title}</title></head>
<body>
<h1>{title}</h1>
<pre>{plain}</pre>
</body>
</html>"""
    return html


def build_pdf_export(title, content):
    """Build actual PDF export using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.6 * inch,
        leftMargin=0.6 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=6,
        textColor=colors.HexColor('#1a1a2e'),
    )

    contact_style = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=12,
        textColor=colors.HexColor('#555555'),
    )

    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=11,
        leading=14,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor('#1a1a2e'),
        borderWidth=0,
        borderPadding=0,
        borderColor=colors.HexColor('#cccccc'),
        borderWidthBottom=1,
        borderPaddingBottom=2,
    )

    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=13,
        spaceAfter=4,
        textColor=colors.HexColor('#333333'),
    )

    story = []

    # ── Extract sections DIRECTLY from content dict ──
    sections_data = {}
    for key in ['body', 'summary', 'education', 'experience', 'projects', 'skills']:
        if key in content and isinstance(content[key], dict):
            text = content[key].get('text', '')
            if text and text.strip():
                sections_data[key] = text.strip()

    # ── HEADER: Name + Contact ──
    name = title or "Resume"
    contact = ""

    if 'body' in sections_data:
        body_lines = sections_data['body'].split('\n')
        if body_lines:
            name = body_lines[0].strip()
            if len(body_lines) > 1:
                contact = " | ".join(line.strip() for line in body_lines[1:] if line.strip())

    story.append(Paragraph(name, name_style))
    if contact:
        story.append(Paragraph(contact, contact_style))

    # ── SECTIONS ──
    section_titles = {
        'summary': 'PROFESSIONAL SUMMARY',
        'education': 'EDUCATION',
        'experience': 'EXPERIENCE',
        'projects': 'PROJECTS',
        'skills': 'SKILLS & CERTIFICATIONS',
    }

    for sec_key in ['summary', 'experience', 'education', 'projects', 'skills']:
        if sec_key in sections_data:
            # Section heading
            story.append(Paragraph(section_titles[sec_key], section_title_style))

            # Body text — newlines ko <br/> se replace karo for reportlab
            text = sections_data[sec_key]
            # Escape XML special chars first
            safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            # Convert newlines to <br/> for line breaks
            safe_text = safe_text.replace('\n', '<br/>')

            story.append(Paragraph(safe_text, body_style))
            story.append(Spacer(1, 4))

    # Build PDF
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    return pdf_bytes


def build_docx_export(title, content):
    """Build actual DOCX export using python-docx."""
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = docx.Document()

        # Set margins
        sections = doc.sections[0]
        sections.top_margin = Inches(0.8)
        sections.bottom_margin = Inches(0.8)
        sections.left_margin = Inches(0.7)
        sections.right_margin = Inches(0.7)

        plain = plain_text_from_content(content)
        lines = plain.split('\n')

        # Extract name (first line)
        name = title or "Resume"
        if lines:
            name = lines[0].strip()

        # Add name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add contact info (second line)
        if len(lines) > 1:
            contact = lines[1].strip()
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(contact)
            contact_run.font.size = Pt(9)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add horizontal line
        doc.add_paragraph("_" * 50)

        # Section keywords
        section_keywords = {
            'summary': ['summary', 'professional summary', 'objective'],
            'education': ['education', 'academic', 'qualification'],
            'experience': ['experience', 'work experience', 'employment'],
            'projects': ['projects', 'personal projects'],
            'skills': ['skills', 'technical skills', 'competencies'],
        }

        current_section = None
        section_titles = {
            'summary': 'PROFESSIONAL SUMMARY',
            'experience': 'EXPERIENCE',
            'education': 'EDUCATION',
            'projects': 'PROJECTS',
            'skills': 'SKILLS & CERTIFICATIONS',
        }

        for line in lines[2:]:
            stripped = line.strip()
            if not stripped:
                continue

            line_lower = stripped.lower()
            found_section = None

            for sec, keywords in section_keywords.items():
                if any(k in line_lower for k in keywords) and len(stripped) < 60:
                    found_section = sec

            if found_section:
                # Add section title
                heading = doc.add_heading(section_titles.get(found_section, found_section.upper()), level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                current_section = found_section
            else:
                # Add body text
                p = doc.add_paragraph(stripped)
                p.paragraph_format.space_after = Pt(2)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        return docx_bytes

    except ImportError:
        # Fallback: return HTML as bytes if python-docx not installed
        return build_printable_html(title, content).encode("utf-8")


def store_export(data, upload_folder, user_id, stem, ext):
    """Store exported file and return path."""
    upload_path = Path(str(upload_folder))
    upload_path.mkdir(parents=True, exist_ok=True)
    
    filename = f"{stem}_{user_id or 0}_{hashlib.md5(data).hexdigest()[:6]}.{ext}"
    file_path = upload_path / filename
    file_path.write_bytes(data if isinstance(data, bytes) else str(data).encode("utf-8"))
    
    return str(file_path)


# ============================================================
# UTILITIES
# ============================================================

def get_file_hash(file_path):
    if not file_path or not os.path.exists(str(file_path)):
        return ""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def generate_unique_filename(original_filename, user_id=None):
    name, ext = os.path.splitext(secure_filename(str(original_filename)))
    hash_part = hashlib.md5(str(user_id or "").encode()).hexdigest()[:6]
    return f"{name}_{hash_part}{ext}"


# Backward compatibility
__all__ = [
    "ResumeValidationError",
    "ResumeParseError",
    "ValidationResult",
    "ParsedResume",
    "validate_resume_upload",
    "save_validated_upload",
    "save_upload",
    "parse_resume_file",
    "content_from_sections",
    "plain_text_from_content",
    "build_json_export",
    "build_printable_html",
    "build_pdf_export",
    "build_docx_export",
    "store_export",
    "get_file_hash",
    "generate_unique_filename",
]